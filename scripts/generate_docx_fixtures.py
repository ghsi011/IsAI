"""Generate synthetic DOCX fixtures for tests.

All content is invented; nothing here comes from a real document. Run:

    uv run python scripts/generate_docx_fixtures.py [output_dir]

Tests call the builder functions directly with tmp_path, so the committed repo
carries no .docx binaries at all.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

LOREM_ACADEMIC = (
    "The retrospective cohort included 412 patients treated between 2015 and 2019 "
    "at three tertiary centers. Median follow-up was 4.2 years, and survival "
    "outcomes were assessed with the Kaplan-Meier method alongside multivariable "
    "Cox regression adjusted for age, stage, and comorbidity burden."
)

FORMULAIC = (
    "Moreover, it is important to note that the findings highlight the importance "
    "of early intervention. Furthermore, these results underscore the significance "
    "of comprehensive assessment. Additionally, the evidence demonstrates the "
    "crucial role of multidisciplinary collaboration in achieving optimal outcomes."
)

HEBREW = "המחקר הראה כי התוצאות היו מובהקות סטטיסטית בקרב קבוצת ההתערבות."

SMART_PUNCTUATION = (
    "As Smith noted, “the effect size (d = 0.42) was modest” — yet the 95% CI "
    "excluded zero… and the α-level of .05 was pre-registered."
)


def build_simple(path: Path) -> Path:
    """Headings, body paragraphs, a list, one table — order-sensitive layout."""
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(LOREM_ACADEMIC)
    doc.add_paragraph("Short paragraph.")
    doc.add_heading("Methods", level=1)
    doc.add_paragraph("Before-table paragraph describing the cohort in detail. " * 3)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Variable"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Patients enrolled"
    table.cell(1, 1).text = "412 across three centers"
    doc.add_paragraph("After-table paragraph interpreting the tabulated values. " * 3)
    doc.add_heading("Statistical analysis", level=2)
    doc.add_paragraph("Analyses used R 4.3 with two-sided tests. " * 4)
    for item in ("First listed criterion", "Second listed criterion"):
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("")  # deliberately empty
    doc.add_paragraph(FORMULAIC)
    doc.save(str(path))
    return path


def build_unicode(path: Path) -> Path:
    doc = Document()
    doc.add_heading("ממצאים", level=1)
    doc.add_paragraph(HEBREW)
    doc.add_paragraph(SMART_PUNCTUATION)
    doc.add_paragraph("Math: ∆H = −41.8 kJ·mol⁻¹ at 25 °C (p ≤ 0.001).")
    doc.save(str(path))
    return path


def build_nested_and_merged(path: Path) -> Path:
    """A table with vertically merged cells and a nested table inside a cell."""
    doc = Document()
    doc.add_paragraph("Paragraph before the complex table.")
    table = doc.add_table(rows=3, cols=2)
    merged = table.cell(0, 0).merge(table.cell(1, 0))
    merged.text = "Merged label"
    table.cell(0, 1).text = "Top right"
    table.cell(1, 1).text = "Middle right"
    table.cell(2, 0).text = "Bottom left"
    inner = table.cell(2, 1).add_table(rows=1, cols=1)
    inner.cell(0, 0).text = "Nested cell content"
    doc.add_paragraph("Paragraph after the complex table.")
    doc.save(str(path))
    return path


def build_thesis(path: Path, *, paragraphs: int = 6) -> Path:
    """A small thesis-like document for end-to-end runs."""
    doc = Document()
    doc.add_heading("A Synthetic Literature Review", level=1)
    doc.add_heading("Background", level=2)
    for i in range(paragraphs):
        body = LOREM_ACADEMIC if i % 2 == 0 else FORMULAIC
        doc.add_paragraph(f"Paragraph {i + 1}. {body}")
    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph(
        "In conclusion, it is important to note that these findings highlight the "
        "importance of the topic and underscore the significance of further research."
    )
    doc.save(str(path))
    return path


def build_with_hyperlink_and_breaks(path: Path) -> Path:
    """Visible hyperlink text and an explicit line break inside one paragraph."""
    doc = Document()
    p = doc.add_paragraph("See the trial registry")
    p_el = p._p  # pyright: ignore[reportPrivateUsage]
    hyperlink = p_el.makeelement(qn("w:hyperlink"), {})
    run = p_el.makeelement(qn("w:r"), {})
    text_el = p_el.makeelement(qn("w:t"), {})
    text_el.text = " ClinicalTrials.gov entry "
    run.append(text_el)
    hyperlink.append(run)
    p_el.append(hyperlink)
    tail_run = p.add_run("for the protocol.")
    tail_run.add_break()
    p.add_run("Second visual line of the same paragraph.")
    doc.save(str(path))
    return path


ALL_BUILDERS = {
    "simple.docx": build_simple,
    "unicode.docx": build_unicode,
    "nested_merged.docx": build_nested_and_merged,
    "thesis.docx": build_thesis,
    "hyperlink_breaks.docx": build_with_hyperlink_and_breaks,
}


def main() -> None:
    out_dir = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("tests/fixtures")
    out_dir.mkdir(parents=True, exist_ok=True)
    for name, builder in ALL_BUILDERS.items():
        print("writing", builder(out_dir / name))


if __name__ == "__main__":
    main()
